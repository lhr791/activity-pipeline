from docx import Document
import sys

doc = Document(r'c:\Users\xtt\Desktop\my-agent-test\2026.2月 竞品所活动 .docx')

with open('docx_content.txt', 'w', encoding='utf-8') as f:
    f.write("=== PARAGRAPHS ===\n")
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text + "\n")

    f.write("\n=== TABLES ===\n")
    for ti, table in enumerate(doc.tables):
        f.write(f"\n--- Table {ti+1} ---\n")
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            f.write(" | ".join(cells) + "\n")

print("Done! Saved to docx_content.txt")
