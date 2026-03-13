"""
交易所活动整合 Word 文档生成器
按照参考文档格式生成：汇总表 + 每个交易所的详细活动文案
"""
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)


def load_events_from_supabase() -> dict:
    """从 Supabase summaries 表读取最新的活动数据。
    返回 dict: active_events, expired_events, version_pairs, events(兼容)"""
    from utils import get_supabase
    db = get_supabase()
    result = (
        db.table("summaries")
        .select("summary")
        .order("created_at", desc=True)
        .limit(1)
        .execute()
    )
    if not result.data:
        print("[ERR] Supabase 中没有找到任何 summary 数据")
        return {"active_events": [], "expired_events": [], "version_pairs": [], "events": []}
    try:
        parsed = json.loads(result.data[0]["summary"])
        # 兼容新旧格式
        if "active_events" in parsed:
            return parsed
        else:
            # 旧格式：只有 events，需要在此做分拣
            from datetime import date as _date
            today = _date.today()
            events = parsed.get("events", [])
            active, expired = [], []
            for e in events:
                end = e.get("end_date", "") or ""
                try:
                    end_dt = datetime.strptime(end[:10], "%Y-%m-%d").date()
                    if end_dt < today:
                        e["status"] = "expired"
                        expired.append(e)
                    else:
                        e["status"] = "active"
                        active.append(e)
                except (ValueError, IndexError):
                    e["status"] = "active"
                    active.append(e)
            # 合并手动录入
            import os as _os
            manual_path = _os.path.join(_os.path.dirname(__file__), "manual_events.json")
            if _os.path.exists(manual_path):
                try:
                    with open(manual_path, "r", encoding="utf-8") as mf:
                        manual = json.load(mf)
                    manual = manual if isinstance(manual, list) else manual.get("events", [])
                    for me in manual:
                        me.setdefault("status", "expired")
                    expired.extend(manual)
                except Exception:
                    pass

            # 版本对比
            from summarizer import find_version_pairs
            pairs = find_version_pairs(active, expired)

            return {
                "active_events": active,
                "expired_events": expired,
                "version_pairs": pairs,
                "events": events,
            }
    except (json.JSONDecodeError, TypeError) as e:
        print(f"[ERR] 解析 summary JSON 失败: {e}")
        return {"active_events": [], "expired_events": [], "version_pairs": [], "events": []}


def _make_tg_link(chat_id: int, message_id: int) -> str:
    """生成 TG 私有频道跳转链接。"""
    internal_id = str(chat_id).replace("-100", "", 1)
    return f"https://t.me/c/{internal_id}/{message_id}"


def enrich_events_with_tg_links(events: list[dict]) -> None:
    """从 raw_messages 表查询最近消息，按 exchange 名匹配注入 source_links。
    过滤汇总帖，优先匹配详细活动内容，按活动时间范围过滤。"""
    from utils import get_supabase
    db = get_supabase()

    result = (
        db.table("raw_messages")
        .select("chat_id,message_id,text,sent_at")
        .order("sent_at", desc=True)
        .limit(1000)
        .execute()
    )
    if not result.data:
        return

    messages = result.data

    # 用于检测汇总帖的交易所名列表
    all_exchange_names = {
        "ourbit", "zoomex", "woox", "fameex", "bydfi", "btcc", "voox",
        "weex", "lbank", "tapbit", "ascendex", "bitmart", "orangex",
        "toobit", "xt ", "deepcoin", "picol", "phemex", "hotcoin",
        "bitrue", "kucoin", "bybit", "bitget",
    }

    def count_exchanges_in_text(text: str) -> int:
        text_lower = text.lower()
        return sum(1 for ex in all_exchange_names if ex in text_lower)

    for ev in events:
        exchange = ev.get("exchange", "").strip()
        if not exchange:
            continue
        if ev.get("source_links"):
            continue

        # 跳过手动录入的活动
        sources = ev.get("sources", [])
        if isinstance(sources, list) and "用户手动录入" in sources:
            continue
        if isinstance(sources, str) and "手动录入" in sources:
            continue

        exchange_lower = exchange.lower()

        # 确定时间窗口：有 end_date 的过期活动只匹配活动期间前后的消息
        end_date_str = ev.get("end_date", "") or ""
        start_date_str = ev.get("start_date", "") or ""
        time_filter_start = None
        time_filter_end = None
        try:
            if end_date_str:
                from datetime import timedelta
                end_dt = datetime.strptime(end_date_str[:10], "%Y-%m-%d")
                time_filter_end = (end_dt + timedelta(days=7)).strftime("%Y-%m-%d")
                if start_date_str:
                    start_dt = datetime.strptime(start_date_str[:10], "%Y-%m-%d")
                    time_filter_start = (start_dt - timedelta(days=7)).strftime("%Y-%m-%d")
                else:
                    time_filter_start = (end_dt - timedelta(days=60)).strftime("%Y-%m-%d")
        except (ValueError, IndexError):
            pass

        candidates = []  # (text_length, link)

        for msg in messages:
            text = msg.get("text") or ""
            text_lower = text.lower()

            if exchange_lower not in text_lower:
                continue

            # 时间过滤
            if time_filter_start and time_filter_end:
                sent = (msg.get("sent_at") or "")[:10]
                if sent < time_filter_start or sent > time_filter_end:
                    continue

            # 跳过汇总帖
            if count_exchanges_in_text(text) >= 3:
                continue

            # 跳过太短的消息
            if len(text) < 100:
                continue

            link = _make_tg_link(msg["chat_id"], msg["message_id"])
            candidates.append((len(text), link))

        # 按消息长度降序排序，取最详细的前 2 条
        candidates.sort(key=lambda x: x[0], reverse=True)
        best_links = [link for _, link in candidates[:2]]

        if best_links:
            ev["source_links"] = best_links


# ── 交易所名称标准化 ──
NAME_MAP = {
    "ourbit": "OurBit", "zoomex": "Zoomex", "woox pro": "WOOX Pro",
    "wooxpro": "WOOX Pro", "fameex": "FameEX", "bydfi": "BYDFI",
    "btcc": "BTCC", "voox": "VOOX", "weex": "WEEX", "lbank": "LBank",
    "tapbit": "Tapbit", "ascendex": "AscendEX", "bitmart": "BitMart",
    "orangex": "OrangeX", "toobit": "Toobit", "xt": "XT",
    "deepcoin": "Deepcoin", "picol": "Picol", "phemex": "Phemex",
    "hotcoin": "Hotcoin", "bitrue": "Bitrue", "kucoin": "KuCoin",
    "woo x": "WOOX Pro", "woox": "WOOX Pro",
}


def normalize_exchange(name: str) -> str:
    return NAME_MAP.get(name.lower().strip(), name)


def add_hyperlink(paragraph, url: str, text: str, font_size=Pt(8), color=RGBColor(0, 0, 238)):
    """在段落中添加可点击的超链接。"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), str(color) if color else "0000EE")
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size.pt * 2))  # half-points
    rPr.append(sz)
    new_run.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    new_run.append(t_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def get_offset_category(loss_offset: int) -> str:
    """根据亏损抵扣比例分类"""
    if loss_offset <= 0:
        return "zero"   # 0%
    elif loss_offset < 100:
        return "partial"  # 33~50%
    else:
        return "full"   # 100%


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_summary_table(doc, events: list[dict]):
    """生成参考文档顶部的汇总表"""
    # 按 loss_offset 分组
    groups = {"zero": [], "partial": [], "full": [], "other": []}

    for ev in events:
        ex = normalize_exchange(ev.get("exchange", ""))
        offset = ev.get("loss_offset", -1)
        commission = ev.get("commission_rate", -1)
        bonus_type = ev.get("bonus_type", "") or ""

        entry = {
            "exchange": ex,
            "bonus_type": bonus_type,
            "offset": offset,
            "commission": commission,
        }

        if offset < 0 or offset is None:
            groups["other"].append(entry)
        elif offset == 0:
            groups["zero"].append(entry)
        elif offset < 100:
            groups["partial"].append(entry)
        else:
            groups["full"].append(entry)

    # 去重：同一交易所只出现一次（取最高返佣）
    for key in groups:
        seen = {}
        for e in groups[key]:
            ex = e["exchange"]
            if ex not in seen or e["commission"] > seen[ex]["commission"]:
                seen[ex] = e
        groups[key] = list(seen.values())

    doc.add_heading("🚀  汇总整合 🚀", level=1)

    # --- 亏损抵扣 0% ---
    if groups["zero"]:
        p = doc.add_paragraph()
        run = p.add_run("═" * 15 + " 亏损抵扣 0% " + "═" * 15)
        run.bold = True
        run.font.size = Pt(10)
        for e in groups["zero"]:
            bt = e["bonus_type"] if e["bonus_type"] else "赠金"
            comm = f"{e['commission']}%返佣" if e["commission"] >= 0 else ""
            p = doc.add_paragraph()
            p.add_run(f"{e['exchange']}   {bt}     {comm}").font.size = Pt(10)

    # --- 亏损抵扣 33~50% ---
    if groups["partial"]:
        p = doc.add_paragraph()
        run = p.add_run("═" * 15 + " 亏损抵扣 33~50% " + "═" * 15)
        run.bold = True
        run.font.size = Pt(10)
        for e in sorted(groups["partial"], key=lambda x: x["offset"]):
            comm = f"{e['commission']}%返佣" if e["commission"] >= 0 else ""
            p = doc.add_paragraph()
            p.add_run(f"{e['exchange']}   {e['offset']}%     {comm}").font.size = Pt(10)

    # --- 亏损抵扣 100% ---
    if groups["full"]:
        p = doc.add_paragraph()
        run = p.add_run("═" * 15 + " 亏损抵扣 100% " + "═" * 15)
        run.bold = True
        run.font.size = Pt(10)
        for e in groups["full"]:
            comm = f"{e['commission']}%返佣" if e["commission"] >= 0 else ""
            p = doc.add_paragraph()
            p.add_run(f"{e['exchange']}       {comm}").font.size = Pt(10)

    # --- 其他类型 ---
    if groups["other"]:
        p = doc.add_paragraph()
        run = p.add_run("═" * 15 + " 其他类型 " + "═" * 15)
        run.bold = True
        run.font.size = Pt(10)
        for e in groups["other"]:
            comm = f"{e['commission']}%返佣" if e["commission"] >= 0 else ""
            p = doc.add_paragraph()
            p.add_run(f"{e['exchange']}         {comm}").font.size = Pt(10)

    doc.add_paragraph()  # 分隔


def add_event_detail(doc, event: dict):
    """按参考文档风格输出单个活动的详细文案"""
    exchange = normalize_exchange(event.get("exchange", ""))
    name = event.get("event_name", "")
    max_reward = event.get("max_reward", -1)
    title_reward = f"最高领 ${max_reward:,}！" if max_reward and max_reward > 0 else ""

    # 活动大标题
    heading = doc.add_heading(level=2)
    run = heading.add_run(f"🚀 【{exchange}】{name}")
    run.font.size = Pt(13)

    # 核心规则摘要
    reward = event.get("reward", "")
    requirements = event.get("requirements", "")

    if reward:
        p = doc.add_paragraph()
        run = p.add_run(f"核心规则： 💰 {reward}")
        run.font.size = Pt(10)

    # 活动时间
    start = event.get("start_date", "")
    end = event.get("end_date", "")
    if start or end:
        time_str = ""
        if start and end:
            time_str = f"{start} – {end}"
        elif start:
            time_str = f"{start} 起"
        elif end:
            time_str = f"截止 {end}"
        p = doc.add_paragraph()
        p.add_run(f"🕗 活动时间： {time_str}").font.size = Pt(10)

    # 亏损抵扣 + 返佣
    offset = event.get("loss_offset", -1)
    commission = event.get("commission_rate", -1)
    info_parts = []
    if offset is not None and offset >= 0:
        info_parts.append(f"亏损抵扣 {offset}%")
    if commission is not None and commission >= 0:
        info_parts.append(f"返佣 {commission}%")

    bonus_type = event.get("bonus_type", "")
    if bonus_type:
        info_parts.append(f"赠金类型: {bonus_type}")

    validity = event.get("bonus_validity_days", -1)
    if validity and validity > 0:
        info_parts.append(f"赠金有效期 {validity}天")

    if info_parts:
        p = doc.add_paragraph()
        p.add_run("✨ " + " | ".join(info_parts)).font.size = Pt(10)

    # 参与条件
    if requirements:
        p = doc.add_paragraph()
        run = p.add_run(f"📋 参与条件：")
        run.bold = True
        run.font.size = Pt(10)
        p = doc.add_paragraph()
        p.add_run(requirements).font.size = Pt(10)

    # 最低入金 / 最高奖励 / 交易量
    details = []
    min_dep = event.get("min_deposit", -1)
    if min_dep and min_dep > 0:
        details.append(f"最低入金: {min_dep:,} USDT")
    if max_reward and max_reward > 0:
        details.append(f"最高奖励: {max_reward:,} USDT")
    target_vol = event.get("target_volume", "")
    if target_vol:
        details.append(f"交易量要求: {target_vol}")
    leverage = event.get("leverage_limit", "")
    if leverage:
        details.append(f"杠杆限制: {leverage}")
    withdrawal = event.get("withdrawal_condition", "")
    if withdrawal:
        details.append(f"提现条件: {withdrawal}")

    if details:
        p = doc.add_paragraph()
        run = p.add_run("💵 奖励详情")
        run.bold = True
        run.font.size = Pt(10)
        for d in details:
            p = doc.add_paragraph()
            p.add_run(f"🔹 {d}").font.size = Pt(10)

    # 新用户/KYC
    new_only = event.get("new_users_only")
    kyc = event.get("kyc_required")
    flags = []
    if new_only:
        flags.append("⚠️ 仅限新用户")
    if kyc:
        flags.append("需要 KYC")
    if flags:
        p = doc.add_paragraph()
        p.add_run(" | ".join(flags)).font.size = Pt(9)

    # 避坑指南
    tips = event.get("tips", "")
    if tips:
        p = doc.add_paragraph()
        run = p.add_run("⚠️ 避坑指南：")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(180, 0, 0)
        p = doc.add_paragraph()
        p.add_run(tips).font.size = Pt(10)

    # 活动链接
    link = event.get("link", "")
    if link:
        p = doc.add_paragraph()
        p.add_run("🔗 活动链接: ").font.size = Pt(9)
        add_hyperlink(p, link, link, font_size=Pt(9))

    # 信息来源
    sources = event.get("sources", [])
    if sources:
        p = doc.add_paragraph()
        run = p.add_run(f"📡 信息来源: {', '.join(sources)}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # TG 原文跳转链接（每个频道仅保留最新的一条）
    source_links = event.get("source_links", [])
    if source_links:
        # 按 chat_id 分组，保留 message_id 最大的
        import re as _re
        best_per_channel: dict[str, tuple[int, str]] = {}
        for link in source_links:
            m = _re.search(r'/c/(\d+)/(\d+)', link)
            if m:
                chat_id, msg_id = m.group(1), int(m.group(2))
                if chat_id not in best_per_channel or msg_id > best_per_channel[chat_id][0]:
                    best_per_channel[chat_id] = (msg_id, link)
            else:
                best_per_channel[link] = (0, link)
        display_links = [v[1] for v in best_per_channel.values()]

        p = doc.add_paragraph()
        run_label = p.add_run("📲 TG原文: ")
        run_label.font.size = Pt(8)
        
        for i, tg_link in enumerate(display_links):
            if i > 0:
                p.add_run("  ").font.size = Pt(8)
            add_hyperlink(p, tg_link, f"[消息{i+1}]", font_size=Pt(8), color=RGBColor(0, 102, 204))

    # 分隔线
    doc.add_paragraph("─" * 50)


def add_expired_section(doc, expired_events: list[dict]):
    """添加『已过期活动』部分"""
    if not expired_events:
        return

    doc.add_page_break()
    heading = doc.add_heading('[已过期] 历史活动', level=1)

    p = doc.add_paragraph()
    run = p.add_run(f"以下 {len(expired_events)} 个活动已过期，仅供参考对比。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 按交易所分组
    ex_groups: dict[str, list[dict]] = {}
    for ev in expired_events:
        ex = normalize_exchange(ev.get("exchange", ""))
        if ex not in ex_groups:
            ex_groups[ex] = []
        ex_groups[ex].append(ev)

    for ex_name in sorted(ex_groups.keys()):
        for ev in ex_groups[ex_name]:
            end = ev.get("end_date", "")
            # 标题带过期标记
            heading = doc.add_heading(level=2)
            run = heading.add_run(
                f"[已过期] 【{normalize_exchange(ev.get('exchange', ''))}】"
                f"{ev.get('event_name', '')}"
            )
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(150, 150, 150)

            if end:
                p = doc.add_paragraph()
                run = p.add_run(f"截止时间: {end}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(180, 0, 0)

            # 简要展示核心信息
            reward = ev.get("reward", "")
            if reward:
                p = doc.add_paragraph()
                p.add_run(f"奖励: {reward}").font.size = Pt(9)

            reqs = ev.get("requirements", "")
            if reqs:
                p = doc.add_paragraph()
                p.add_run(f"条件: {reqs}").font.size = Pt(9)

            offset = ev.get("loss_offset", -1)
            if offset is not None and offset >= 0:
                p = doc.add_paragraph()
                p.add_run(f"亏损抵扣: {offset}%").font.size = Pt(9)

            # 信息来源
            sources = ev.get("sources", [])
            if sources:
                p = doc.add_paragraph()
                run = p.add_run(f"📡 信息来源: {', '.join(sources)}")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(128, 128, 128)

            # TG 原文跳转链接
            source_links = ev.get("source_links", [])
            if source_links:
                p = doc.add_paragraph()
                p.add_run("📲 TG原文: ").font.size = Pt(8)
                for i, tg_link in enumerate(source_links):
                    if i > 0:
                        p.add_run("  ").font.size = Pt(8)
                    add_hyperlink(p, tg_link, f"[消息{i+1}]", font_size=Pt(8), color=RGBColor(0, 102, 204))

            doc.add_paragraph("---" * 20)


def add_version_comparison(doc, version_pairs: list[dict]):
    """添加『版本变更对比』部分：高亮入金/交易量/奖励差异"""
    if not version_pairs:
        return

    doc.add_page_break()
    doc.add_heading('活动版本变更对比', level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        f"以下 {len(version_pairs)} 组活动检测到版本更新，"
        "对比入金条件、交易量、奖励三个核心字段的变化。"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    compare_fields = [
        ("reward", "奖励"),
        ("requirements", "参与条件"),
        ("target_volume", "交易量要求"),
        ("min_deposit", "最低入金"),
        ("max_reward", "最高奖励"),
        ("loss_offset", "亏损抵扣"),
        ("commission_rate", "返佣比例"),
    ]

    for pair in version_pairs:
        old_ev = pair["old"]
        new_ev = pair["new"]
        exchange = normalize_exchange(new_ev.get("exchange", ""))

        heading = doc.add_heading(level=2)
        heading.add_run(
            f"【{exchange}】{new_ev.get('event_name', '')}"
        ).font.size = Pt(12)

        p = doc.add_paragraph()
        old_end = old_ev.get('end_date', '?')
        new_end = new_ev.get('end_date', '?')
        run = p.add_run(f"旧版截止: {old_end}  |  新版截止: {new_end}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # 逐字段对比
        for field_key, field_label in compare_fields:
            old_val = str(old_ev.get(field_key, "") or "")
            new_val = str(new_ev.get(field_key, "") or "")
            if old_val == new_val:
                continue
            if not old_val and not new_val:
                continue

            p = doc.add_paragraph()
            run = p.add_run(f">> {field_label}:")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 0, 0)

            p = doc.add_paragraph()
            run = p.add_run(f"  旧: {old_val[:200]}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(150, 150, 150)

            p = doc.add_paragraph()
            run = p.add_run(f"  新: {new_val[:200]}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 100, 0)

        doc.add_paragraph("---" * 20)


def generate_word(
    active_events: list,
    output_path: str,
    expired_events: list | None = None,
    version_pairs: list | None = None,
):
    """生成完整的活动报告 Word 文档（四段结构）"""
    expired_events = expired_events or []
    version_pairs = version_pairs or []

    # 从 raw_messages 反查 TG 跳转链接
    enrich_events_with_tg_links(active_events)
    enrich_events_with_tg_links(expired_events)

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 标题
    now = datetime.now()
    month_str = f"{now.year}.{now.month}月"
    title = doc.add_heading(f'{month_str} 竞品所活动', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 生成日期 + 统计
    total = len(active_events) + len(expired_events)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"生成日期: {now.strftime('%Y-%m-%d %H:%M')} | "
        f"当期 {len(active_events)} 个 + 过期 {len(expired_events)} 个 | "
        f"数据来源: TG 活动频道监听 + AI 整合"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

    # 第一部分：汇总表（仅当期活动）
    add_summary_table(doc, active_events)

    # 第二部分：当期活动详细
    exchange_events: dict[str, list[dict]] = {}
    for ev in active_events:
        ex = normalize_exchange(ev.get("exchange", ""))
        if ex not in exchange_events:
            exchange_events[ex] = []
        exchange_events[ex].append(ev)

    for ex_name in sorted(exchange_events.keys()):
        for ev in exchange_events[ex_name]:
            add_event_detail(doc, ev)

    # 第三部分：过期活动
    add_expired_section(doc, expired_events)

    # 第四部分：版本对比
    add_version_comparison(doc, version_pairs)

    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        f"-- 报告结束 | "
        f"当期 {len(active_events)} 个活动，过期 {len(expired_events)} 个，"
        f"{len(version_pairs)} 组版本变更 --"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"[OK] Word 文档已生成: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) > 1 and not sys.argv[1].startswith("--"):
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            data = json.load(f)
        if isinstance(data, list):
            active = data
            expired = []
            pairs = []
        else:
            active = data.get("active_events", data.get("events", []))
            expired = data.get("expired_events", [])
            pairs = data.get("version_pairs", [])
    else:
        print("从 Supabase 读取最新活动数据...")
        result = load_events_from_supabase()
        active = result.get("active_events", [])
        expired = result.get("expired_events", [])
        pairs = result.get("version_pairs", [])

    if not active and not expired:
        print("[ERR] 没有活动数据")
        sys.exit(1)

    print(f"读取到 {len(active)} 个当期 + {len(expired)} 个过期活动")

    output_dir = os.path.join(os.path.dirname(__file__), "output")
    os.makedirs(output_dir, exist_ok=True)
    now = datetime.now()
    date_str = now.strftime("%Y%m%d_%H%M")
    month_str = f"{now.year}.{now.month}月"
    default_output = os.path.join(output_dir, f"{month_str} 竞品所活动_{date_str}.docx")
    output = sys.argv[2] if len(sys.argv) > 2 else default_output

    generate_word(active, output, expired_events=expired, version_pairs=pairs)
