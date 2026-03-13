"""
交易所赠金规则独立提取工具
从活动链接抓取页面 → AI 提取结构化赠金规则 → Word 报告

Usage:
    python3 bonus_rules.py --urls URL1 URL2 ...
    python3 bonus_rules.py --exchanges Bitrue Phemex OurBit
    python3 bonus_rules.py --from-db
    python3 bonus_rules.py --from-db --dry     # 只抓取，不调 AI
    python3 bonus_rules.py --bonus-rules       # 从活动页面自动发现赠金说明链接
"""

import argparse
import json
import os
import re
import time
from datetime import datetime
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from utils import get_supabase, get_openai, logger
from enrich_events import (
    fetch_page, ALLOWED_DOMAINS, SKIP_PATH_PATTERNS,
)

# 赠金说明页常在帮助中心域名，fetch_page 需要这些域名
SUPPORT_DOMAINS = {
    "support.bitrue.com", "bitmart.zendesk.com",
    "www.kucoin.com", "phemex.com",
}

db = get_supabase()
ai = get_openai()

# ── AI Prompt ──────────────────────────────────────────────────────────────

RULES_PROMPT = """你是加密货币交易所活动数据提取专家。

从给定的活动页面内容中提取完整的赠金规则。

## 提取目标（严格 JSON）
{
  "exchange": "交易所名",
  "event_name": "活动名称",
  "tiers": [
    {"deposit": 1000, "volume": 500000, "bonus": 200, "note": "可选备注"}
  ],
  "min_deposit": 1000,
  "max_reward": 5000,
  "target_volume": 32000000,
  "loss_offset_pct": 100,
  "bonus_validity_days": 7,
  "commission_rate": "Maker 0.02% / Taker 0.06%",
  "new_users_only": false,
  "withdrawal_kills_bonus": true,
  "key_restrictions": ["限前100名", "10天内不能提现"],
  "summary": "一句话总结这个活动"
}

## 规则
- tiers 按入金金额从小到大排列
- 数值用数字不用字符串
- 如果页面有多个活动，分别提取，用 JSON 数组返回
- 页面无法识别为活动页面时返回 null
- 所有输出用中文
"""

# 赠金机制规则提取 prompt（帮助中心文章 → 结构化整理）
BONUS_DESC_PROMPT = """你是加密货币交易所赠金规则分析专家。

这是一个交易所帮助中心的赠金使用说明页面。请将赠金机制规则整理成结构化格式。

## 重要：只提取赠金机制规则
✅ 要：赠金是什么、怎么用、亏损/费用怎么扣、过期怎么回收、盈利如何处理
❌ 不要：入金多少送多少、活动档位、提币比例等活动特定内容

## 输出格式（严格 JSON）
{
  "exchange": "交易所名",
  "bonus_name": "赠金类型名称",
  "bonus_type": "trading_bonus / voucher / trial_fund / coupon",
  "sections": [
    {
      "title": "什么是合约赠金",
      "content": "保留原文关键描述，不要压缩"
    },
    {
      "title": "发放形式",
      "content": "1. 直接发放到合约账户\n2. 卡券形式..."
    },
    {
      "title": "使用规则",
      "content": "完整的使用规则列表"
    },
    {
      "title": "亏损与费用扣除",
      "content": "扣除机制说明",
      "examples": ["范例1: ...", "范例2: ..."]
    },
    {
      "title": "过期与回收",
      "content": "回收规则"
    },
    {
      "title": "注意事项",
      "content": "重要提醒"
    }
  ]
}

## 规则
- sections 的 title 和 content 必须用中文
- content 要详细，保留原文的关键信息和数字，不要压缩成一句话
- examples 字段可选，有范例时一定要保留
- 如果页面讲了多种赠金类型（如 bonus + voucher），分别提取，用 JSON 数组返回
- 无法识别时返回 null
"""

# AI 精筛链接 prompt
LINK_FILTER_PROMPT = """以下是从加密货币交易所活动页面提取的超链接。
请判断哪些是"赠金/优惠券使用规则说明页"的链接。

✅ 算赠金说明页（选它）：
- 帮助中心/FAQ/Support 中解释赠金(bonus)怎么用、怎么解锁、怎么提现的页面
- 优惠券(coupon/voucher)使用规则说明页
- 体验金(trial fund)使用说明页
- 标题含 "bonus rules"、"coupon"、"voucher"、"trial fund" 的 support 文章

❌ 不算（排除）：
- 活动招募页、注册页、保险基金页、借贷页、官方验证页
- 交易页面、Earn/理财页面、邀请好友页

只返回符合的链接编号，用 JSON 数组，如 [1, 3]。
如果没有符合的，返回 []。"""


def extract_rules(url: str, page_content: str, prompt: str = None) -> list[dict] | None:
    """AI 提取赠金规则。"""
    system_prompt = prompt or RULES_PROMPT
    resp = ai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"页面 URL: {url}\n\n{page_content}"},
        ],
        temperature=0.1,
        max_tokens=2048,
    )
    raw = resp.choices[0].message.content.strip()

    if raw.lower() in ("null", "none"):
        return None

    raw = re.sub(r"^```(?:json)?\s*\n?", "", raw)
    raw = re.sub(r"\n?```\s*$", "", raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r'[\[{].*[}\]]', raw, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group())
            except json.JSONDecodeError:
                logger.warning("Failed to parse AI response for %s", url)
                return None
        else:
            return None

    if isinstance(data, dict):
        data = [data]
    return data


# ── 赠金说明链接发现 ──────────────────────────────────────────────────────────

BONUS_LINK_KEYWORDS = [
    "bonus", "trial", "fund", "voucher", "coupon",
    "赠金", "体验金", "试用金", "奖金",
]


def discover_bonus_rule_links() -> list[dict]:
    """从数据库活动页面中自动发现赠金说明超链接。"""
    from playwright.sync_api import sync_playwright

    # 读取所有活动链接
    r = db.table("summaries").select("summary").order("created_at", desc=True).limit(1).execute()
    if not r.data:
        return []
    s = json.loads(r.data[0]["summary"])

    targets = []
    seen_urls = set()
    for ev in s.get("events", []):
        link = (ev.get("link") or "").strip()
        if not link or link in seen_urls:
            continue
        if any(p in link.lower() for p in SKIP_PATH_PATTERNS):
            continue
        seen_urls.add(link)
        targets.append({"exchange": ev.get("exchange", "?"), "url": link})

    logger.info("共 %d 个活动页面待扫描", len(targets))

    found = {}  # url -> {exchange, text, href}

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)

        for i, t in enumerate(targets):
            ex = t["exchange"]
            url = t["url"]
            logger.info("[%d/%d] %s | %s", i + 1, len(targets), ex, url[:70])

            try:
                page = browser.new_page()
                page.goto(url, wait_until="networkidle", timeout=12000)

                # 尝试展开详情区域
                for sel in ["text=Event details", "text=Details", "text=Rules",
                            "text=Terms", "text=详情", "text=规则", "text=条款"]:
                    try:
                        page.click(sel, timeout=1500)
                        page.wait_for_timeout(800)
                    except:
                        pass

                page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                page.wait_for_timeout(500)

                links = page.evaluate("""() => {
                    return Array.from(document.querySelectorAll('a[href]')).map(a => ({
                        href: a.href,
                        text: (a.innerText || a.textContent || '').trim().substring(0, 150)
                    })).filter(l => l.href && !l.href.startsWith('javascript'))
                }""")

                page.close()

                for l in links:
                    href_l = l["href"].lower()
                    text_l = l["text"].lower()
                    if any(kw in href_l or kw in text_l for kw in BONUS_LINK_KEYWORDS):
                        # 排除：导航菜单、注册页、同页锚点
                        if ("/register" in href_l or "/signup" in href_l
                                or "/invite" in href_l or href_l.endswith("#")
                                or "protect" in href_l or "shield" in href_l
                                or "my-tasks" in href_l or "reward-center" in href_l
                                or "funding-rate" in href_l or "download" in href_l
                                or "invite" in href_l):
                            continue
                        if not l["text"] or len(l["text"]) < 3:
                            continue
                        if l["href"] not in found:
                            found[l["href"]] = {
                                "exchange": ex,
                                "text": l["text"],
                                "href": l["href"],
                            }
                            logger.info("  🎯 [%s] → %s", l["text"][:50], l["href"])

            except Exception as e:
                logger.warning("  ❌ %s: %s", ex, str(e)[:80])

        browser.close()

    candidates = list(found.values())
    logger.info("关键词预筛: %d 个候选链接", len(candidates))

    if not candidates:
        return []

    # AI 精筛
    link_list = "\n".join(
        f"{i+1}. [{c['text'][:80]}] → {c['href']}" for i, c in enumerate(candidates)
    )
    resp = ai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": LINK_FILTER_PROMPT},
            {"role": "user", "content": link_list},
        ],
        temperature=0,
        max_tokens=200,
    )
    raw = resp.choices[0].message.content.strip()
    try:
        selected = json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r'\[.*?\]', raw)
        selected = json.loads(match.group()) if match else []

    results = [candidates[i - 1] for i in selected if 1 <= i <= len(candidates)]
    logger.info("AI 精筛: %d/%d 个确认为赠金说明链接", len(results), len(candidates))
    for r in results:
        logger.info("  ✅ [%s] → %s", r["text"][:50], r["href"])
    return results


# ── 链接获取 ────────────────────────────────────────────────────────────────

def urls_from_db() -> list[dict]:
    """从数据库 active_events 读取活动链接。"""
    r = db.table("summaries").select("summary").order("created_at", desc=True).limit(1).execute()
    if not r.data:
        return []
    s = json.loads(r.data[0]["summary"])
    results = []
    for ev in s.get("active_events", []):
        link = (ev.get("link") or "").strip()
        if link and not any(p in link.lower() for p in SKIP_PATH_PATTERNS):
            parsed = urlparse(link)
            if parsed.hostname in ALLOWED_DOMAINS:
                results.append({
                    "url": link,
                    "exchange": ev.get("exchange", "?"),
                    "event_name": ev.get("event_name", "?"),
                })
    return results


def urls_from_search(exchanges: list[str]) -> list[dict]:
    """Google 搜索交易所活动链接。"""
    results = []
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    for exchange in exchanges:
        query = f'"{exchange}" deposit bonus activity site:{exchange.lower()}.com'
        logger.info("Searching: %s", query)
        try:
            resp = requests.get(
                "https://www.google.com/search",
                params={"q": query, "num": 5},
                headers=headers,
                timeout=10,
            )
            soup = BeautifulSoup(resp.text, "html.parser")
            for a in soup.find_all("a", href=True):
                href = a["href"]
                # Google 结果链接格式
                if "/url?q=" in href:
                    url = href.split("/url?q=")[1].split("&")[0]
                    parsed = urlparse(url)
                    if parsed.hostname in ALLOWED_DOMAINS:
                        if not any(p in url.lower() for p in SKIP_PATH_PATTERNS):
                            results.append({
                                "url": url,
                                "exchange": exchange,
                                "event_name": "搜索结果",
                            })
                            logger.info("  Found: %s", url)
        except Exception as e:
            logger.warning("Search failed for %s: %s", exchange, e)
        time.sleep(1)

    # 去重
    seen = set()
    unique = []
    for r in results:
        if r["url"] not in seen:
            seen.add(r["url"])
            unique.append(r)
    return unique


# ── Word 输出 ───────────────────────────────────────────────────────────────

def generate_word_report(all_rules: list[dict], output_path: str, mode: str = "activity"):
    """Word 报告。mode='bonus_desc' 用赠金机制规则模板。"""
    doc = Document()

    if mode == "bonus_desc":
        _generate_bonus_desc_report(doc, all_rules)
    else:
        _generate_activity_report(doc, all_rules)

    doc.save(output_path)
    logger.info("Word 报告已保存: %s", output_path)


def _generate_bonus_desc_report(doc, all_rules: list[dict]):
    """赠金机制规则报告。"""
    doc.add_heading("竞品交易所赠金机制规则汇总", level=0)
    doc.add_paragraph(f"共整理 {len(all_rules)} 个交易所赠金规则")

    for idx, rule in enumerate(all_rules, 1):
        exchange = rule.get("exchange", "?")
        bonus_name = rule.get("bonus_name", rule.get("event_name", "?"))
        bonus_type = rule.get("bonus_type", "")
        source_url = rule.get("activity_link", "")

        # 交易所标题
        doc.add_heading(f"{idx}. {exchange} — {bonus_name}", level=1)

        # 类型 + 来源链接
        meta_parts = []
        if bonus_type:
            meta_parts.append(f"类型: {bonus_type}")
        if source_url:
            meta_parts.append(f"来源: {source_url}")
        if meta_parts:
            p = doc.add_paragraph()
            run = p.add_run(" | ".join(meta_parts))
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 200)

        # 按 section 展示详细规则
        sections = rule.get("sections", [])
        for sec in sections:
            title = sec.get("title", "")
            content = sec.get("content", "")
            examples = sec.get("examples", [])

            if title:
                doc.add_heading(title, level=2)

            if content:
                for line in content.split("\n"):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)

            if examples:
                for ex in examples:
                    p = doc.add_paragraph()
                    run = p.add_run(ex)
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(80, 80, 80)

        if idx < len(all_rules):
            doc.add_page_break()


def _generate_activity_report(doc, all_rules: list[dict]):
    """活动规则报告（原模板）。"""
    # 标题
    title = doc.add_heading("交易所赠金规则对比报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # 汇总表
    doc.add_heading("汇总", level=1)
    summary_table = doc.add_table(rows=1, cols=7)
    summary_table.style = "Light Grid Accent 1"
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["交易所", "活动名", "最低入金", "最高赠金", "亏损抵扣", "有效期", "限新用户"]
    for i, h in enumerate(headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    for rule in all_rules:
        row = summary_table.add_row()
        vals = [
            rule.get("exchange", "?"),
            rule.get("event_name", "?"),
            f"${rule.get('min_deposit', '?')}",
            f"${rule.get('max_reward', '?')}",
            f"{rule.get('loss_offset_pct', '?')}%",
            f"{rule.get('bonus_validity_days', '?')}天",
            "是" if rule.get("new_users_only") else "否",
        ]
        for i, v in enumerate(vals):
            row.cells[i].text = str(v)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_page_break()

    # 详情
    for idx, rule in enumerate(all_rules, 1):
        exchange = rule.get("exchange", "?")
        event_name = rule.get("event_name", "?")

        doc.add_heading(f"{idx}. {exchange} — {event_name}", level=1)

        # 一句话总结
        summary_text = rule.get("summary", "")
        if summary_text:
            p = doc.add_paragraph()
            run = p.add_run(summary_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

        # 档位表
        tiers = rule.get("tiers", [])
        if tiers:
            doc.add_heading("奖励档位", level=2)
            tier_table = doc.add_table(rows=1, cols=4)
            tier_table.style = "Light Grid Accent 1"
            tier_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            tier_headers = ["档位", "入金 (USDT)", "交易量 (USDT)", "赠金 (USDT)"]
            for i, h in enumerate(tier_headers):
                cell = tier_table.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)

            for ti, tier in enumerate(tiers, 1):
                row = tier_table.add_row()
                dep = tier.get("deposit", "-")
                vol = tier.get("volume", "-")
                bonus = tier.get("bonus", "-")
                note = tier.get("note", "")

                # 格式化大数字
                def fmt_num(n):
                    if n == "-" or n is None:
                        return "-"
                    n = int(n)
                    if n >= 1_000_000:
                        return f"{n/1_000_000:.1f}M"
                    elif n >= 1_000:
                        return f"{n/1_000:.0f}K"
                    return str(n)

                vals = [str(ti), fmt_num(dep), fmt_num(vol), fmt_num(bonus)]
                for i, v in enumerate(vals):
                    row.cells[i].text = v
                    for p in row.cells[i].paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

        # 关键参数
        doc.add_heading("关键参数", level=2)
        params = [
            ("最低入金", f"${rule.get('min_deposit', '?')} USDT"),
            ("最高赠金", f"${rule.get('max_reward', '?')} USDT"),
            ("最高档交易量", f"${rule.get('target_volume', '?')} USDT" if rule.get('target_volume') else "无要求"),
            ("亏损抵扣", f"{rule.get('loss_offset_pct', '?')}%"),
            ("赠金有效期", f"{rule.get('bonus_validity_days', '?')} 天"),
            ("手续费率", rule.get("commission_rate", "未知")),
            ("限新用户", "是" if rule.get("new_users_only") else "否"),
            ("提现取消赠金", "是" if rule.get("withdrawal_kills_bonus") else "否"),
        ]
        param_table = doc.add_table(rows=0, cols=2)
        param_table.style = "Light Grid Accent 1"
        for label, value in params:
            row = param_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

        # 关键限制
        restrictions = rule.get("key_restrictions", [])
        if restrictions:
            doc.add_heading("关键限制 / 避坑", level=2)
            for r_text in restrictions:
                doc.add_paragraph(r_text, style="List Bullet")

        # 活动链接
        link = rule.get("activity_link", "")
        if link:
            p = doc.add_paragraph()
            run = p.add_run(f"活动链接: {link}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 200)

        if idx < len(all_rules):
            doc.add_page_break()

    doc.save(output_path)
    logger.info("Word 报告已保存: %s", output_path)


# ── 主流程 ──────────────────────────────────────────────────────────────────

def run(args):
    # 1. 获取链接
    targets = []
    if args.bonus_rules:
        # 从活动页面自动发现赠金说明链接
        bonus_links = discover_bonus_rule_links()
        for bl in bonus_links:
            targets.append({"url": bl["href"], "exchange": bl["exchange"], "event_name": bl["text"]})
        # 赠金说明页用专用 prompt
        args._use_bonus_desc_prompt = True
    elif args.urls:
        for url in args.urls:
            parsed = urlparse(url)
            exchange = parsed.hostname or "Unknown"
            exchange = exchange.replace("www.", "").split(".")[0].capitalize()
            targets.append({"url": url, "exchange": exchange, "event_name": "手动输入"})
    elif args.from_db:
        targets = urls_from_db()
    elif args.exchanges:
        targets = urls_from_search(args.exchanges)
    else:
        logger.error("请指定 --bonus-rules, --urls, --from-db, 或 --exchanges")
        return

    logger.info("共 %d 个目标链接", len(targets))

    if not targets:
        logger.error("未找到任何有效链接")
        return

    # 2. 抓取 + AI 提取
    all_rules = []
    for i, target in enumerate(targets):
        url = target["url"]
        logger.info("[%d/%d] %s - %s", i + 1, len(targets), target["exchange"], url)

        content = fetch_page(url)
        # 帮助中心域名 fetch_page 可能不认识，额外尝试
        if not content:
            parsed = urlparse(url)
            if parsed.hostname in SUPPORT_DOMAINS:
                try:
                    resp = requests.get(url, timeout=10, verify=False,
                                       headers={"User-Agent": "Mozilla/5.0"})
                    if resp.status_code == 200 and len(resp.text) > 200:
                        soup = BeautifulSoup(resp.text, "html.parser")
                        content = soup.get_text(separator="\n", strip=True)[:8000]
                except Exception:
                    pass
        if not content:
            logger.warning("  ✗ 无法抓取: %s", url)
            continue
        logger.info("  ✓ 抓取 %d 字符", len(content))

        if args.dry:
            logger.info("  [DRY] 跳过 AI")
            continue

        # 选择 prompt
        use_prompt = BONUS_DESC_PROMPT if getattr(args, '_use_bonus_desc_prompt', False) else RULES_PROMPT

        rules = extract_rules(url, content, prompt=use_prompt)
        if rules:
            for rule in rules:
                rule["activity_link"] = url
                if not rule.get("exchange"):
                    rule["exchange"] = target["exchange"]
            all_rules.extend(rules)
            logger.info("  ✓ 提取 %d 条赠金规则", len(rules))
        else:
            logger.info("  - 无法识别")

        time.sleep(1)

    if args.dry:
        logger.info("DRY RUN 完成")
        return

    logger.info("共提取 %d 个活动规则", len(all_rules))

    if not all_rules:
        logger.error("未提取到任何规则")
        return

    # 3. 输出
    # JSON
    os.makedirs("output", exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    json_path = f"output/赠金规则_{ts}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(all_rules, f, ensure_ascii=False, indent=2)
    logger.info("JSON 已保存: %s", json_path)

    # Word
    word_path = f"output/赠金规则_{ts}.docx"
    report_mode = "bonus_desc" if getattr(args, '_use_bonus_desc_prompt', False) else "activity"
    generate_word_report(all_rules, word_path, mode=report_mode)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="交易所赠金规则提取工具")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--bonus-rules", action="store_true", help="从活动页面自动发现赠金说明链接")
    group.add_argument("--urls", nargs="+", help="直接指定活动 URL")
    group.add_argument("--from-db", action="store_true", help="从数据库读取活动链接")
    group.add_argument("--exchanges", nargs="+", help="按交易所名搜索活动链接")
    parser.add_argument("--dry", action="store_true", help="只抓取页面，不调 AI")
    args = parser.parse_args()
    run(args)
