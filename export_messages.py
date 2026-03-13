"""
导出 Supabase raw_messages 为按频道分组的 Word 文档。
不做 AI 整合，仅原文导出。
"""
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from utils import get_supabase

db = get_supabase()

CHANNEL_NAMES = {
    -1002000478651: "Coinscalper Channel",
    -1003061431387: "증정금 No1 레드터틀 채널",
    -1003601907317: "Dalchuni Crypto Events",
    -1002770517188: "Global Loha(Crypto Event)",
    -1003500837149: "Redturtle_Global_Events",
    -1003389965115: "Global Exchange Event Summary",
}


def export() -> str:
    """导出原始消息为 docx，返回文件路径。"""
    result = (
        db.table("raw_messages")
        .select("chat_id, message_id, text, sent_at")
        .order("sent_at", desc=True)
        .limit(1000)
        .execute()
    )
    messages = result.data
    print(f"共 {len(messages)} 条消息")

    # 按频道分组
    by_channel: dict[int, list] = {}
    for m in messages:
        cid = m["chat_id"]
        if cid not in by_channel:
            by_channel[cid] = []
        by_channel[cid].append(m)

    # 生成 Word 文档
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(9)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    now = datetime.now()
    title = doc.add_heading(f'TG 原始聊天记录', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"导出时间: {now.strftime('%Y-%m-%d %H:%M')} | 共 {len(messages)} 条消息，{len(by_channel)} 个频道")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    for chat_id, msgs in sorted(by_channel.items()):
        channel_name = CHANNEL_NAMES.get(chat_id, str(chat_id))

        doc.add_heading(f"📺 {channel_name}", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"Chat ID: {chat_id} | {len(msgs)} 条消息")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # 按时间正序
        msgs.sort(key=lambda x: x["sent_at"])
        for m in msgs:
            ts = m["sent_at"][:16].replace("T", " ")
            p = doc.add_paragraph()
            run = p.add_run(f"[{ts}]")
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(80, 80, 80)

            p = doc.add_paragraph()
            p.add_run(m["text"]).font.size = Pt(9)

            # 分隔线
            p = doc.add_paragraph()
            run = p.add_run("─" * 50)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(200, 200, 200)

    # 保存
    output_dir = os.path.join(os.path.dirname(__file__), "output")
    os.makedirs(output_dir, exist_ok=True)
    date_str = now.strftime("%Y%m%d_%H%M")
    output_path = os.path.join(output_dir, f"TG原始聊天记录_{date_str}.docx")
    doc.save(output_path)
    print(f"[OK] 原始聊天记录导出: {output_path}")
    return output_path


if __name__ == "__main__":
    export()
